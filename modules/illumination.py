import streamlit as st
import time
from utils.gemini_ai import get_soul_rebel_consultant
from utils.supabase_db import save_brand_data, load_brand_data

def run(user_id):
    st.title("✨ Scriptly Labs: Master Voice Blueprint")
    st.caption("Forging your high-fidelity brand authority engine: Subsection by Subsection.")
    st.write("---")
    
    # 1. INITIALIZE MASTER DATA & UI STATE
    db_data = load_brand_data(user_id)
    brand_data = db_data if db_data else {}

    # The Cognitive Blueprint: Explicitly mapping the exact weight, length, and rhythm of the document
    guide_structure = [
        {
            "id": "identity_big_idea",
            "label": "SECTION 1: BRAND IDENTITY",
            "sub": "Big Idea & Meaning",
            "blueprint": "- Structural Target: A single, rousing, evocative affirmation phrase followed by a deep narrative dive spanning 7-8 expansive paragraphs.\n- Tone: Poetic, soul-stirring, conversational yet authoritative."
        },
        {
            "id": "identity_vision_mission",
            "label": "SECTION 1: BRAND IDENTITY",
            "sub": "Vision & Mission",
            "blueprint": "- Structural Target: 2 distinct sub-blocks featuring punchy, bold core declarations followed by a detailed, vivid 3-4 paragraph 'What it Means' narrative expansion.\n- Tone: Visionary, bold, revolutionary."
        },
        {
            "id": "process_methodology",
            "label": "SECTION 1: BRAND IDENTITY",
            "sub": "Our Transformation Process",
            "blueprint": "- Structural Target: A high-level introduction paragraph establishing a 3-way systemic change, followed by 3 robust, multi-paragraph conceptual deep-dives.\n- Tone: Strategic, defiant, highly structured, and action-oriented."
        },
        {
            "id": "anchors_culture_values",
            "label": "SECTION 1: BRAND IDENTITY",
            "sub": "Core Anchors: Culture & Values",
            "blueprint": "- Structural Target: 1 introductory narrative paragraph followed by 7 distinct, highly detailed value pillars (e.g., Humility, Diversity, Curiosity).\n- Tone: Humble yet fiercely confident, deeply relational, and 'hella smart'."
        },
        {
            "id": "anchors_beliefs_behaviours",
            "label": "SECTION 1: BRAND IDENTITY",
            "sub": "Core Anchors: Beliefs & Behaviours",
            "blueprint": "- Structural Target: 2 comprehensive, bulleted manifestos. 'Beliefs' section strictly begins with 'We believe...' (10-12 items). 'Behaviours' section strictly uses 'We always...' statements (6-8 items).\n- Tone: Unwavering, non-negotiable, and programmatic."
        },
        {
            "id": "positioning_1soul",
            "label": "SECTION 2: BRAND POSITIONING",
            "sub": "Primary Anchor Statement & Meaning",
            "blueprint": "- Structural Target: A singular, definitive corporate anchor sentence followed by a powerful 5-6 paragraph 'What it Means' narrative.\n- Tone: Masterful, elite, authoritative, and uncompromising."
        },
        {
            "id": "positioning_offerings",
            "label": "SECTION 2: BRAND POSITIONING",
            "sub": "Our Offering Matrix",
            "blueprint": "- Structural Target: 3 comprehensive, separate ecosystem deep-dives mapping out exact value propositions for Staff, Clients, and Communities.\n- Tone: Generous, expert, professional, and impact-driven."
        },
        {
            "id": "positioning_audience",
            "label": "SECTION 2: BRAND POSITIONING",
            "sub": "Target Audience Profiles",
            "blueprint": "- Structural Target: 2 exhaustive, distinct profile breakdowns detailing 'Our Allies' and 'Our Target Partners/Staff', followed by a visceral list defining industry 'Harm'.\n- Tone: Empathetic, descriptive, sharp, and raw."
        },
        {
            "id": "expression_slogan",
            "label": "SECTION 3: BRAND EXPRESSION",
            "sub": "Strategic Slogan & Manifest",
            "blueprint": "- Structural Target: A crisp, evocative core corporate slogan backed by a 3-paragraph conceptual manifesto defining what that slogan delivers.\n- Tone: Inspiring, cinematic, premium, and visually descriptive."
        },
        {
            "id": "expression_voice_personality",
            "label": "SECTION 3: BRAND EXPRESSION",
            "sub": "Brand Voice & Personality Archetype",
            "blueprint": "- Structural Target: 2 comprehensive segments detailing the verbal identity guidelines ('Human, compelling, and hella smart') and the core character archetype.\n- Tone: Highly articulate, commanding, and profoundly loving."
        },
        {
            "id": "ties_legacy",
            "label": "SECTION 4: CORE LEVERAGE TIES",
            "sub": "Brand Legacy",
            "blueprint": "- Structural Target: 1 high-level 'living legacy' statement paragraph, followed by 6 distinct, structured 'We want to be known for...' declarations.\n- Tone: Historic, monumental, ancestral, and enduring."
        },
        {
            "id": "ties_markers",
            "label": "SECTION 4: CORE LEVERAGE TIES",
            "sub": "Key Voice Performance Indicators (KPIs)",
            "blueprint": "- Structural Target: An exclusive framework consisting of 6-8 razor-sharp metrics written strictly from the perspective of the practitioner as a 'Do I...' or 'Am I...' self-audit question.\n- Tone: Introspective, challenging, deeply accountable, and uncompromising."
        }
    ]

    # --- PERSISTENT DATA INITIALIZATION & SMART RESUME ENGINE ---
    if "guide_idx" not in st.session_state:
        st.session_state.guide_idx = 0

    if "guide_content" not in st.session_state or not st.session_state.guide_content:
        st.session_state.guide_content = {}
        highest_completed_idx = -1
        
        # Scan the database rows to discover what subsections have already been finalized
        for i, item in enumerate(guide_structure):
            col_name = f"guide_part_{item['id']}"
            if col_name in brand_data and brand_data[col_name]:
                st.session_state.guide_content[item['id']] = brand_data[col_name]
                highest_completed_idx = i
        
        # Guard loop navigation fallback redirects to sync state safely
        if st.session_state.get("current_nav", "2. The Soul Guide") == "2. The Soul Guide":
            st.session_state.guide_idx = highest_completed_idx + 1
                
    if "rev" not in st.session_state:
        st.session_state.rev = 0

    curr_idx = st.session_state.guide_idx

    # --- FINAL MASTER COMPILATION (WITH DB SAVE & ADVANCE) ---
    if curr_idx >= len(guide_structure):
        st.success("🎉 All blueprint sectors meticulously forged! Your Master Voice Guide is complete.")
        
        # Stitch the entire narrative together for the screen view
        full_text = "\n\n".join([st.session_state.guide_content.get(item['id'], '') for item in guide_structure])
        st.text_area("Scriptly Labs: Master Blueprint Final Asset", value=full_text, height=450)
        
        # Action Column Layout for the Final Gate
        c1, c2 = st.columns(2)
        
        with c1:
            if st.button("💾 Save Final Strategy & Advance", use_container_width=True, type="primary"):
                with st.spinner("Locking Master Voice Blueprint into the database..."):
                    save_brand_data(user_id, full_text, chamber="soul_guide")
                    
                    try:
                        from utils.supabase_db import get_supabase_client
                        supabase = get_supabase_client()
                        
                        # 1. Update navigation state in profile layers
                        supabase.table("profiles").update({"last_nav": "3. Brand Guardian"}).eq("user_id", user_id).execute()
                        
                        # 2. Assign routing key parameters back to central controller matrix
                        st.session_state.target_page = "3. Brand Guardian"
                        
                    except Exception as e:
                        st.warning("Strategy saved locally, but navigation routing advance encountered an obstacle.")
                    
                    st.success("Strategy locked! Routing to Brand Guardian Control...")
                    time.sleep(1.0)
                    st.rerun()
                    
        with c2:
            # Word Document Export
            try:
                from docx import Document
                from io import BytesIO
                
                docx_buffer = BytesIO()
                doc = Document()
                
                doc.add_heading("SCRIPTLY LABS: MASTER VOICE BLUEPRINT GUIDE", level=0)
                doc.add_paragraph("Generated via Scriptly Labs Engine • High-Fidelity Content Architecture")
                doc.add_page_break()
                
                for item in guide_structure:
                    section_text = st.session_state.guide_content.get(item['id'], '').strip()
                    if section_text:
                        if "🚨 BLUEPRINT MATRIX INQUIRY" in section_text:
                            section_text = section_text.split("🚨 BLUEPRINT MATRIX INQUIRY")[0].strip()
                            
                        doc.add_heading(item['label'], level=1)
                        doc.add_heading(item['sub'], level=2)
                        
                        for line in section_text.split("\n"):
                            clean_line = line.strip()
                            if not clean_line:
                                continue
                            if clean_line.startswith("-") or clean_line.startswith("*"):
                                doc.add_paragraph(clean_line[1:].strip(), style='List Bullet')
                            else:
                                doc.add_paragraph(clean_line)
                        doc.add_paragraph("\n")
                
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                st.download_button(
                    label="📄 Export to Word Document (.docx)",
                    data=docx_buffer,
                    file_name="Master_Voice_Blueprint_Final.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except ImportError:
                st.error("Please add python-docx to your software environments to enable Word export functionality.")

        st.write("---")
        if st.button("⬅️ Go Back to Edit Sections", use_container_width=True):
            st.session_state.guide_idx = len(guide_structure) - 1
            st.rerun()
        return

    curr_sub = guide_structure[curr_idx]
    sub_key = curr_sub["id"]

    # 2. WORKSPACE LAYOUT
    col1, col2 = st.columns([3, 2])

    with col1:
        st.subheader(f"{curr_sub['label']}")
        st.info(f"**Forging Blueprint Element:** {curr_sub['sub']}")
        
        # A. INTUITIVE GENERATION (Using the deep blueprint + Soul Audit data)
        if sub_key not in st.session_state.guide_content:
            with st.spinner(f"Scriptly Engine analyzing and drafting {curr_sub['sub']}..."):
                methodology = f"""
                ROLE: Expert Creative Director & Brand Architect at Scriptly Labs.
                TASK: Synthesize the raw discovery audit inputs perfectly into the designated layout module: '{curr_sub['sub']}'.
                
                STRICT STRUCTURAL AND TONAL BLUEPRINT:
                {curr_sub['blueprint']}
                
                VOICE STANDARD: You are an elite professional corporate copywriter and voice strategist who outputs deeply compelling, human, authoritative, and 'hella smart' layout blocks.
                
                GAP CAPTURE: If vital source components are absent, draft a robust strategic baseline baseline framework matching the blueprint criteria, then append a clean section titled '🚨 BLUEPRINT MATRIX INQUIRY' at the absolute bottom containing 3 precise, targeted refinement questions.
                """
                draft = get_soul_rebel_consultant(f"Draft {curr_sub['sub']}", methodology + "\n\nDISCOVERY SOURCE FILE DATA:\n" + str(brand_data))
                st.session_state.guide_content[sub_key] = draft

        # B. ANTI-SHADOWING WORKSPACE CONTAINER
        workspace_key = f"area_{sub_key}_v{st.session_state.rev}"
        current_text = st.session_state.guide_content[sub_key]
        
        if "🚨 BLUEPRINT MATRIX INQUIRY" in current_text:
            st.warning("The Scriptly Architect requires missing layout details to satisfy the target framework depth.")

        edited_text = st.text_area(
            "Refine Strategic Voice Narrative Content:", 
            value=current_text, 
            height=450, 
            key=workspace_key
        )
        st.session_state.guide_content[sub_key] = edited_text

        # C. ADAPTIVE COLLABORATIVE SYNTHESIS ENGINE
        st.write("---")
        user_input = st.chat_input(f"Feed raw operational details or reply to matrix inquiries for {curr_sub['sub']}...")
        
        if user_input:
            with st.spinner("Braiding and synthesizing inputs into master blueprint layers..."):
                update_prompt = f"""
                ROLE: Expert Creative Director & Brand Architect at Scriptly Labs.
                TASK: Elevate the user's raw real-world insights into our premium corporate voice, and surgically braid the updates into the existing layout draft. Expand the section blocks—do not compress or truncate. Maintain absolute adherence to blueprint constraints: {curr_sub['blueprint']}
                
                If the user insight resolves a query block inside '🚨 BLUEPRINT MATRIX INQUIRY', remove that query chunk entirely from the output.
                Return ONLY the complete, newly generated and expanded modular text layout block.
                
                EXISTING DRAFT TO REVISE:
                {edited_text}
                """
                new_narrative = get_soul_rebel_consultant(user_input, update_prompt)
                st.session_state.guide_content[sub_key] = new_narrative
                st.session_state.rev += 1
                st.rerun()

        # D. NAVIGATION BUTTONS IN WORKSPACE
        st.write("---")
        b_col1, b_col2 = st.columns(2)
        with b_col1:
            if curr_idx > 0:
                if st.button("⬅️ Save & Go Back", use_container_width=True):
                    save_brand_data(user_id, edited_text, chamber=f"guide_part_{sub_key}")
                    st.session_state.guide_idx -= 1
                    st.session_state.rev = 0
                    st.rerun()
        with b_col2:
            is_light = "🚨 BLUEPRINT MATRIX INQUIRY" in edited_text
            if st.button("🔥 Commit & Advance Section", use_container_width=True, disabled=is_light):
                save_brand_data(user_id, edited_text, chamber=f"guide_part_{sub_key}")
                
                # If we are finishing up the last index step, compile the master copy block to the DB
                if curr_idx == len(guide_structure) - 1:
                    full_compiled_text = "\n\n".join([st.session_state.guide_content.get(item['id'], '') for item in guide_structure])
                    save_brand_data(user_id, full_compiled_text, chamber="soul_guide")
                    
                st.session_state.guide_idx += 1
                st.session_state.rev = 0
                st.rerun()

    # 3. ROADMAP & INTERACTIVE JUMP PANEL (RIGHT COLUMN)
    with col2:
        st.subheader("📋 Blueprint Roadmap Track")
        st.caption("Select any cleared index element to audit or adapt previous copies.")
        st.write("---")
        
        for i, item in enumerate(guide_structure):
            if i < curr_idx:
                if st.button(f"✅ {item['sub']}", key=f"jump_{item['id']}", use_container_width=True, help="Revisit this cleared sector"):
                    save_brand_data(user_id, edited_text, chamber=f"guide_part_{sub_key}")
                    st.session_state.guide_idx = i
                    st.session_state.rev = 0
                    st.rerun()
            elif i == curr_idx:
                st.markdown(f"**👉 {item['sub']}** (Active Composition)")
            else:
                st.caption(f"⚪ {item['sub']}")
        
        st.write("---")
        if is_light:
            st.error("Compliance Alert: This sector cannot be committed until all questions inside '🚨 BLUEPRINT MATRIX INQUIRY' are systematically answered and removed.")